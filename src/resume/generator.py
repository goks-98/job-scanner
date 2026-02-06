"""
Resume generator to create ATS-optimized .docx files.
"""

import os
from datetime import datetime
from typing import Dict, Any, List
from pathlib import Path
import logging

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

from src.utils.helpers import get_project_root


class ResumeGenerator:
    """Generate ATS-optimized .docx resumes."""
    
    def __init__(self, output_dir: str = None):
        """
        Initialize the resume generator.
        
        Args:
            output_dir: Directory to save generated resumes
        """
        self.logger = logging.getLogger(self.__class__.__name__)
        
        if output_dir is None:
            output_dir = get_project_root() / "data" / "resumes"
        
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate(self, optimized_resume: Dict[str, Any], job_id: str, 
                job_title: str, company: str) -> str:
        """
        Generate a .docx resume file.
        
        Args:
            optimized_resume: Optimized resume data from ResumeOptimizer
            job_id: Unique job identifier
            job_title: Target job title
            company: Target company name
            
        Returns:
            Path to generated .docx file
        """
        self.logger.info(f"Generating resume for {job_title} at {company}")
        
        doc = Document()
        
        # Set up document styles
        self._setup_styles(doc)
        
        # Add sections
        self._add_header(doc, optimized_resume)
        self._add_summary(doc, optimized_resume.get('summary', ''))
        self._add_skills(doc, optimized_resume.get('skills', {}))
        self._add_experience(doc, optimized_resume.get('experience', []))
        self._add_certifications(doc, optimized_resume.get('certifications', []))
        self._add_education(doc, optimized_resume.get('education', []))
        
        # Generate filename
        safe_company = "".join(c for c in company if c.isalnum() or c in ' -_')[:30]
        safe_title = "".join(c for c in job_title if c.isalnum() or c in ' -_')[:30]
        timestamp = datetime.now().strftime('%Y%m%d')
        filename = f"{optimized_resume.get('name', 'Resume').replace(' ', '_')}_{safe_company}_{timestamp}.docx"
        
        output_path = self.output_dir / filename
        doc.save(str(output_path))
        
        self.logger.info(f"Resume saved to: {output_path}")
        return str(output_path)
    
    def _setup_styles(self, doc: Document):
        """Set up ATS-friendly document styles."""
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'  # ATS-friendly font
        font.size = Pt(11)
        font.color.rgb = RGBColor(0, 0, 0)
        
        # Paragraph formatting
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Heading styles
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = 'Calibri'
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.styles['Heading 1'].font.size = Pt(16)
        doc.styles['Heading 2'].font.size = Pt(12)
    
    def _add_header(self, doc: Document, resume: Dict[str, Any]):
        """Add resume header with contact information."""
        # Name
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(resume.get('name', 'Your Name'))
        name_run.bold = True
        name_run.font.size = Pt(18)
        
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(resume.get('target_title', 'Data Engineer'))
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = RGBColor(70, 70, 70)
        
        # Contact info on one line
        contact = resume.get('contact', {})
        contact_parts = []
        
        if contact.get('phone'):
            contact_parts.append(contact['phone'])
        if contact.get('email'):
            contact_parts.append(contact['email'])
        if contact.get('location'):
            contact_parts.append(contact['location'])
        
        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.add_run(' | '.join(contact_parts))
        
        # LinkedIn on separate line
        if contact.get('linkedin'):
            linkedin_para = doc.add_paragraph()
            linkedin_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            linkedin_para.add_run(contact['linkedin'])
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_section_heading(self, doc: Document, title: str):
        """Add a section heading with underline."""
        heading = doc.add_paragraph()
        heading_run = heading.add_run(title.upper())
        heading_run.bold = True
        heading_run.font.size = Pt(11)
        
        # Add a line under heading
        border_para = doc.add_paragraph()
        border_para.paragraph_format.space_before = Pt(0)
        border_para.paragraph_format.space_after = Pt(6)
    
    def _add_summary(self, doc: Document, summary: str):
        """Add professional summary section."""
        if not summary:
            return
        
        self._add_section_heading(doc, 'Professional Summary')
        
        summary_para = doc.add_paragraph()
        summary_para.add_run(summary)
    
    def _add_skills(self, doc: Document, skills: Dict[str, List[str]]):
        """Add skills section."""
        all_skills = []
        for category, skill_list in skills.items():
            all_skills.extend(skill_list)
        
        if not all_skills:
            return
        
        self._add_section_heading(doc, 'Technical Skills')
        
        # Group skills by category
        if skills.get('technical'):
            tech_para = doc.add_paragraph()
            tech_para.add_run('Programming: ').bold = True
            tech_para.add_run(', '.join(skills['technical']))
        
        if skills.get('tools'):
            tools_para = doc.add_paragraph()
            tools_para.add_run('Tools & Platforms: ').bold = True
            tools_para.add_run(', '.join(skills['tools']))
        
        # If no categorization, just list all
        if not skills.get('technical') and not skills.get('tools'):
            skills_para = doc.add_paragraph()
            skills_para.add_run(', '.join(all_skills))
    
    def _add_experience(self, doc: Document, experience: List[Dict]):
        """Add work experience section."""
        if not experience:
            return
        
        self._add_section_heading(doc, 'Professional Experience')
        
        for exp in experience:
            # Company and title
            exp_header = doc.add_paragraph()
            
            company = exp.get('company', '')
            title = exp.get('title', '')
            dates = exp.get('dates', '')
            location = exp.get('location', '')
            
            # Title and company on same line
            if title:
                exp_header.add_run(title).bold = True
            if company:
                if title:
                    exp_header.add_run(' | ')
                exp_header.add_run(company)
            
            # Location and dates
            if location or dates:
                detail_para = doc.add_paragraph()
                detail_parts = []
                if location:
                    detail_parts.append(location)
                if dates:
                    detail_parts.append(dates)
                detail_run = detail_para.add_run(' | '.join(detail_parts))
                detail_run.font.color.rgb = RGBColor(100, 100, 100)
                detail_run.font.size = Pt(10)
            
            # Bullet points
            for bullet in exp.get('bullets', []):
                bullet_para = doc.add_paragraph(style='List Bullet')
                bullet_para.add_run(bullet)
            
            # Add spacing between jobs
            doc.add_paragraph()
    
    def _add_certifications(self, doc: Document, certifications: List[str]):
        """Add certifications section."""
        if not certifications:
            return
        
        self._add_section_heading(doc, 'Certifications')
        
        for cert in certifications:
            cert_para = doc.add_paragraph(style='List Bullet')
            cert_para.add_run(cert)
    
    def _add_education(self, doc: Document, education: List[Dict]):
        """Add education section."""
        if not education:
            return
        
        self._add_section_heading(doc, 'Education')
        
        for edu in education:
            edu_para = doc.add_paragraph()
            
            degree = edu.get('degree', '')
            institution = edu.get('institution', '')
            dates = edu.get('dates', '')
            
            if degree:
                edu_para.add_run(degree).bold = True
            if institution:
                if degree:
                    edu_para.add_run(' | ')
                edu_para.add_run(institution)
            if dates:
                edu_para.add_run(f' ({dates})')
    
    def generate_sample(self) -> str:
        """Generate a sample resume for testing."""
        sample_data = {
            'name': 'Gokul Subramanian',
            'target_title': 'Senior Data Engineer',
            'contact': {
                'email': 'gokul12cool@gmail.com',
                'phone': '+971 54 302 4015',
                'location': 'Dubai, UAE',
                'linkedin': 'linkedin.com/in/gokul-subramanian-068381191'
            },
            'summary': 'Data Engineer with 5+ years of experience designing and building scalable data platforms. Expert in Python, Spark, and cloud technologies.',
            'skills': {
                'technical': ['Python', 'PySpark', 'SQL', 'Scala'],
                'tools': ['Databricks', 'Snowflake', 'BigQuery', 'Airflow', 'Terraform', 'Docker']
            },
            'experience': [
                {
                    'company': 'Deriv',
                    'title': 'Data Engineer',
                    'location': 'Dubai, UAE',
                    'dates': 'Apr 2025 - Present',
                    'bullets': [
                        'Developed streaming pipelines to ingest millions of trading events daily',
                        'Led migration of event-driven pipelines from GCP Cloud Functions Gen1 to Gen2',
                        'Deployed containerized Python ETL jobs orchestrated with Airflow'
                    ]
                }
            ],
            'certifications': ['Azure Data Engineer Associate', 'Databricks Certified Data Engineer Associate'],
            'education': [{'degree': 'B.E. Aeronautical Engineering', 'institution': 'Madras Institute of Technology'}]
        }
        
        return self.generate(sample_data, 'sample', 'Data Engineer', 'Sample Company')
